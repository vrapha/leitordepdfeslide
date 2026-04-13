"""
DOCX Service — Extração de códigos de questões de arquivos Word (.docx).

Estrutura esperada do Word:
  PARTE 1 — BLUEPRINT (ignorado)
  PARTE 2 — SIMULADO
    Questão N  (N = número)
    BANCA ANO | Questão X do banco original
    [enunciado]
    A. / B. / C. / D. / E.
  QUESTÕES DE RESERVA
    Questão R1, Questão R2, ... (mesmo formato)

Premissas:
  - TODAS as questões estão no Manager → busca agressiva, sem desistência
  - Quantidade variável → detectada automaticamente
  - Resultado em ordem, separado em principais e reservas
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Dict, List, Optional, Tuple

from docx import Document as DocxDocument

from services.pdf_service import (
    QuestionBlock,
    STORAGE_STATE,
    SessionExpiredError,
    wait_for_new_session,
    auto_relogin,
    _ensure_logged_in_and_save_state,
    build_queries_from_enunciado,
    goto_filter_page,
    wait_results,
    parse_listagem_texto,
    validate_question_match,
    SiteQuestion,
    MatchResult,
    count_pdf_alternatives,
)

# ─── Parâmetros agressivos para Word ────────────────────────────────────────
DOCX_MAX_QUERIES   = 120
DOCX_MAX_PAGES     = 8
DOCX_SPECIALTY_IDX = 3
# ────────────────────────────────────────────────────────────────────────────


@dataclass
class DocxQuestion:
    """Questão do Word com indicação se é reserva."""
    block: QuestionBlock
    is_reserva: bool
    label: str   # "Q1", "Q2" ... ou "R1", "R2" ...


# ─────────── Parsing do Word ───────────

def _compact(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip()


def _parse_questao_header(text: str) -> Optional[Tuple[int, bool]]:
    """
    Retorna (numero, is_reserva) se for header de questão, senão None.
    Aceita: 'Questão 1', 'Questão R1', 'Questão R10 | PSU-MG 2025 | ...'
    """
    t = text.strip()
    # Reserva: Questão R1, Questão R2... (pode ter texto extra após o número)
    m = re.match(r"^Quest[aã]o\s+[Rr](\d+)(\s*$|\s*[\|,])", t, re.I)
    if m:
        return int(m.group(1)), True
    # Principal: Questão 1, Questão 2... (pode ter texto extra após o número)
    m = re.match(r"^Quest[aã]o\s+(\d+)(\s*$|\s*[\|,])", t, re.I)
    if m:
        return int(m.group(1)), False
    return None


def _is_banca_line(text: str) -> bool:
    # "banco original" (formato padrão) — o formato "Banco: Q102" vem no header e é tratado pelo |
    return bool(re.search(r"banco\s+original", text, re.I))


def _parse_alternative(text: str) -> Tuple[Optional[str], Optional[str]]:
    m = re.match(r"^([A-E])[\.\)]\s+(.+)", text.strip())
    if m:
        return m.group(1), _compact(m.group(2))
    return None, None


def _is_reserva_section(text: str) -> bool:
    """Detecta marcador de início da seção de questões de reserva."""
    return bool(re.search(r"quest[õo]es?\s+de\s+reserva|reserva[s]?\s*$", text, re.I))


def _is_preamble_line(text: str) -> bool:
    """Linhas que devem ser ignoradas no modo sem-header (títulos, seções, etc.)."""
    if re.search(r"quest[õo]es?\s+selecionadas", text, re.I):
        return True
    if re.search(r"PARTE\s+[123]", text, re.I):
        return True
    if re.search(r"simulado|blueprint|especialidade", text, re.I) and len(text) < 80:
        return True
    return False


def _iter_all_paragraphs(doc):
    """
    Itera todos os parágrafos do documento em ordem, incluindo os que estão
    dentro de células de tabela (que doc.paragraphs ignora).
    """
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.table import Table as DocxTable

    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            yield DocxParagraph(child, doc)
        elif tag == "tbl":
            tbl = DocxTable(child, doc)
            for row in tbl.rows:
                # Usa set para evitar células mescladas duplicadas
                seen = set()
                for cell in row.cells:
                    cid = id(cell._tc)
                    if cid in seen:
                        continue
                    seen.add(cid)
                    for para in cell.paragraphs:
                        yield para


def parse_questoes_from_docx(
    docx_path: str,
    logger: Callable = print,
) -> List[DocxQuestion]:
    """
    Lê .docx e retorna lista de DocxQuestion em ordem (principais primeiro, depois reservas).
    Suporta dois formatos:
      - Com headers: 'Questão 1', 'Questão R1', etc.
      - Sem headers: enunciado direto → A/B/C/D/E → próximo enunciado
    """
    doc = DocxDocument(docx_path)
    # Usa iterador que inclui parágrafos dentro de tabelas (headers em table cells)
    paragraphs = [_compact(p.text) for p in _iter_all_paragraphs(doc)]

    # Localiza início da PARTE 2
    parte2_idx = None
    for i, text in enumerate(paragraphs):
        if re.search(r"PARTE\s+2", text, re.I):
            parte2_idx = i
            break
    if parte2_idx is None:
        for i, text in enumerate(paragraphs):
            if _parse_questao_header(text) is not None:
                parte2_idx = i
                break
    if parte2_idx is None:
        raise RuntimeError(
            "Não foi possível localizar a seção de questões. "
            "Verifique se o arquivo contém 'PARTE 2' ou 'Questão 1'."
        )

    # Detecta se o arquivo usa headers explícitos (formato padrão).
    # Headers reais SEMPRE contêm '|' (Questão N | BANCA | Banco: QXX).
    # "Questão N" sem '|' é referência do banco — ignorada.
    body = [t for t in paragraphs[parte2_idx + 1:] if t]
    has_headers = any(
        _parse_questao_header(t) is not None and '|' in t
        for t in body[:30]
    )

    logger("Detectando questões automaticamente...")

    if has_headers:
        questions = _parse_com_headers(body, logger)
    else:
        logger("Formato sem headers detectado — usando detecção por alternativas.")
        questions = _parse_sem_headers(body, logger)

    principais = [q for q in questions if not q.is_reserva]
    reservas   = [q for q in questions if q.is_reserva]
    logger(f"{len(principais)} questões principais + {len(reservas)} reservas = {len(questions)} total.")
    return questions


def _parse_com_headers(body: List[str], logger: Callable) -> List[DocxQuestion]:
    """Parser original: usa headers 'Questão N' / 'Questão R1'."""
    questions: List[DocxQuestion] = []
    current_numero: Optional[int] = None
    current_is_reserva: bool = False
    current_enunciado_parts: List[str] = []
    current_alternativas: Dict[str, str] = {}
    in_enunciado = False

    def flush():
        nonlocal current_numero, current_is_reserva
        nonlocal current_enunciado_parts, current_alternativas, in_enunciado
        if current_numero is None:
            return
        enunciado = _compact(" ".join(current_enunciado_parts))
        alts = dict(current_alternativas)
        label = f"R{current_numero}" if current_is_reserva else f"Q{current_numero}"
        if len(enunciado.split()) >= 3:
            block = QuestionBlock(
                numero=current_numero,
                tipo="ACESSO_DIRETO",
                enunciado=enunciado,
                alternativas=alts,
                texto_completo=enunciado,
            )
            questions.append(DocxQuestion(block=block, is_reserva=current_is_reserva, label=label))
        current_numero = None
        current_enunciado_parts = []
        current_alternativas = {}
        in_enunciado = False

    for text in body:
        parsed = _parse_questao_header(text)
        if parsed is not None:
            # ── Headers reais SEMPRE têm '|'. Sem '|' = referência do banco. ──
            # Ex: "Questão 83" numa célula adjacente da tabela — ignorar sempre.
            if '|' not in text:
                continue
            flush()
            current_numero, current_is_reserva = parsed
            in_enunciado = True   # header tem | → banca na mesma linha → próximo é enunciado
            continue

        if current_numero is None:
            continue

        if not in_enunciado and _is_banca_line(text):
            in_enunciado = True
            continue

        if not in_enunciado and not current_enunciado_parts and re.match(r"^[A-Z].*\d{4}.*\|", text):
            in_enunciado = True
            continue

        letra, alt_text = _parse_alternative(text)
        if letra:
            # ── FIX 2: alternativas "órfãs" antes do enunciado ───────────────
            # Em layouts de duas colunas, D e E da questão anterior aparecem
            # numa célula ao lado do header atual, ANTES do enunciado real.
            # Regra: só desliga in_enunciado (e sinaliza fim do enunciado) ao
            # encontrar a alternativa "A" — que é sempre a primeira da série
            # real. Letras B-E antes de "A" são tratadas como órfãs: adicionadas
            # mas sem desligar in_enunciado.
            if letra == 'A' and current_enunciado_parts:
                # Alternativa A real encontrada após enunciado coletado.
                # Limpa alternativas órfãs (D, E, etc. coletadas antes do enunciado).
                current_alternativas = {}
                in_enunciado = False
            elif letra != 'A':
                # Pode ser órfã ou pode ser alternativa real (após a A já
                # processada). Só desliga in_enunciado se já coletamos enunciado
                # E já temos a alternativa A (ou seja, estamos na série real).
                if current_enunciado_parts and 'A' in current_alternativas:
                    in_enunciado = False
                # caso contrário, mantém in_enunciado para continuar coletando
            if alt_text:
                current_alternativas[letra] = alt_text
            continue

        if in_enunciado:
            if text in ("[Texto não disponível]", "[Esta questão contém imagem/tabela"):
                current_enunciado_parts.append("questão com imagem sem texto disponível")
            else:
                current_enunciado_parts.append(text)
        elif not current_enunciado_parts:
            # Texto chegou antes de in_enunciado ser True (ex: banca em coluna
            # adjacente lida fora de ordem). Religa a coleta de enunciado.
            if not re.match(r"^quest[aã]o\s+\d+", text, re.I):  # não é header
                in_enunciado = True
                current_enunciado_parts.append(text)

    flush()
    return questions


def _parse_sem_headers(body: List[str], logger: Callable) -> List[DocxQuestion]:
    """
    Parser alternativo para arquivos sem headers 'Questão N'.
    Detecta fronteiras de questão pela alternativa 'A.' — quando aparece após
    ter coletado enunciado, encerra a questão anterior ao finalizar as alternativas.
    Seção de reservas detectada por marcador textual.
    """
    questions: List[DocxQuestion] = []
    q_num = 0
    r_num = 0
    is_reserva = False

    enunciado_parts: List[str] = []
    alternativas: Dict[str, str] = {}
    in_alternatives = False   # True enquanto coletando A/B/C/D/E
    last_letra: Optional[str] = None

    def flush():
        nonlocal q_num, r_num
        enunciado = _compact(" ".join(enunciado_parts))
        alts = dict(alternativas)
        if len(enunciado.split()) < 3:
            return
        if is_reserva:
            r_num += 1
            label = f"R{r_num}"
            numero = r_num
        else:
            q_num += 1
            label = f"Q{q_num}"
            numero = q_num
        block = QuestionBlock(
            numero=numero,
            tipo="ACESSO_DIRETO",
            enunciado=enunciado,
            alternativas=alts,
            texto_completo=enunciado,
        )
        questions.append(DocxQuestion(block=block, is_reserva=is_reserva, label=label))
        enunciado_parts.clear()
        alternativas.clear()

    for text in body:
        # Detecta seção de reservas
        if _is_reserva_section(text):
            # Finaliza última questão principal antes de entrar em reservas
            if in_alternatives and alternativas:
                flush()
                in_alternatives = False
            elif enunciado_parts and not in_alternatives:
                pass  # enunciado sem alternativas — ignora
            is_reserva = True
            continue

        # Ignora linhas de preamble/título
        if _is_preamble_line(text):
            continue

        letra, alt_text = _parse_alternative(text)

        if letra == "A":
            # Início das alternativas → fecha enunciado coletado até aqui
            in_alternatives = True
            if alt_text:
                alternativas[letra] = alt_text
            last_letra = letra
            continue

        if letra and in_alternatives:
            if alt_text:
                alternativas[letra] = alt_text
            last_letra = letra
            continue

        # Texto não é alternativa
        if in_alternatives:
            # Saímos das alternativas → flush da questão completa
            flush()
            in_alternatives = False
            last_letra = None
            # Este texto é início do próximo enunciado
            if text not in ("[Texto não disponível]", "[Esta questão contém imagem/tabela"):
                enunciado_parts.append(text)
            else:
                enunciado_parts.append("questão com imagem sem texto disponível")
        else:
            # Ainda coletando enunciado
            if text not in ("[Texto não disponível]", "[Esta questão contém imagem/tabela"):
                enunciado_parts.append(text)
            else:
                enunciado_parts.append("questão com imagem sem texto disponível")

    # Flush final (última questão pode terminar no fim do arquivo)
    if in_alternatives and alternativas:
        flush()
    elif enunciado_parts and not in_alternatives and not alternativas:
        pass  # enunciado solto sem alternativas — ignora

    return questions


# ─────────── Busca agressiva para Word ───────────

def _find_code_docx(
    page,
    questao: QuestionBlock,
    logger: Callable = print,
    job_id: Optional[str] = None,
) -> Optional[MatchResult]:
    """Busca agressiva — todas as questões estão no Manager."""
    queries = build_queries_from_enunciado(questao.enunciado, questao.alternativas)
    total_alts = count_pdf_alternatives(questao.alternativas) or 4
    seen_codes: set[str] = set()
    best_media: Optional[tuple] = None
    best_baixa: Optional[tuple] = None
    query_count = 0
    max_score_seen = 0

    def _cancelled() -> bool:
        if not job_id:
            return False
        from services.job_manager import get_job as _gj
        j = _gj(job_id)
        return j is not None and j.cancelled

    logger(f"  {len(queries)} queries geradas.")

    for q in queries[:DOCX_MAX_QUERIES]:
        if _cancelled():
            return None
        query_count += 1

        for pnum in range(1, DOCX_MAX_PAGES + 1):
            goto_filter_page(page, q, pnum)
            wait_results(page)

            rows = page.evaluate(
                f"""() => {{
                    const out = [];
                    for (const tr of document.querySelectorAll('table tbody tr')) {{
                        const tds = tr.querySelectorAll('td');
                        if (!tds || tds.length < {DOCX_SPECIALTY_IDX + 1}) continue;
                        const code = (tds[1]?.innerText || '').trim();
                        const desc = (tds[2]?.innerText || '').trim();
                        const esp  = (tds[{DOCX_SPECIALTY_IDX}]?.innerText || '').trim();
                        if (code && desc) out.push({{code, desc, esp}});
                    }}
                    return out;
                }}"""
            )

            if not rows:
                break

            for r in rows[:25]:
                code = (r.get("code") or "").strip()
                if not code or code in seen_codes:
                    continue
                seen_codes.add(code)
                raw_desc = (r.get("desc") or "").strip()
                esp = (r.get("esp") or "").strip()
                if not raw_desc:
                    continue

                enun, alts, is_ad = parse_listagem_texto(raw_desc)
                if not enun:
                    continue

                sq = SiteQuestion(code, enun, alts, is_ad, esp)
                match_ok, score, num_alt = validate_question_match(questao, sq)
                max_score_seen = max(max_score_seen, score)

                if not match_ok:
                    continue

                ratio = num_alt / total_alts if total_alts else 0
                if score >= 90 and ratio >= 0.75:
                    confianca = "ALTA"
                elif score >= 80 and ratio >= 0.60:
                    confianca = "MEDIA"
                else:
                    confianca = "BAIXA"

                result = MatchResult(code, score, num_alt, confianca, is_ad, esp)
                rank = score * 10 + num_alt

                if confianca == "ALTA":
                    return result
                if confianca == "MEDIA":
                    if best_media is None or rank > best_media[1]:
                        best_media = (result, rank)
                else:
                    if best_baixa is None or rank > best_baixa[1]:
                        best_baixa = (result, rank)

            if best_media and best_media[0].score_enunciado >= 90:
                return best_media[0]

    result = (best_media[0] if best_media else None) or (best_baixa[0] if best_baixa else None)
    if not result:
        logger(f"  Não encontrado após {query_count} queries. Melhor score: {max_score_seen}%")
    return result


# ─────────── Extração principal ───────────

def run_docx_extraction(
    docx_path: str,
    logger: Callable = print,
    job_id: Optional[str] = None,
) -> Dict:
    """
    Extrai códigos de TODAS as questões do .docx.
    Retorna dict com 'codes' (lista completa), 'principais' e 'reservas' separados.
    """
    from playwright.sync_api import sync_playwright

    if not Path(docx_path).exists():
        raise FileNotFoundError(f"Arquivo não encontrado: {docx_path}")

    logger("=" * 50)
    logger("EXTRAÇÃO DE QUESTÕES DO WORD")
    logger("=" * 50)

    dq_list = parse_questoes_from_docx(docx_path, logger)
    total = len(dq_list)

    if not dq_list:
        raise RuntimeError("Nenhuma questão foi encontrada no documento.")

    # Se não há sessão salva ou ela expirou, tenta login automático antes de falhar
    if not Path(STORAGE_STATE).exists():
        logger("⚠️  Sessão não encontrada — tentando login automático...")
        if not auto_relogin(logger):
            raise RuntimeError(
                "Sessão do painel não encontrada e login automático falhou. "
                "Configure MANAGER_EMAIL e MANAGER_PASSWORD no Railway."
            )

    principais: List[Dict] = []
    reservas: List[Dict] = []
    found_count = 0

    with sync_playwright() as p:

        def _init_browser():
            b = p.chromium.launch(headless=True, slow_mo=30)
            # Recarrega o storage_state do disco (pode ter sido renovado pelo auto_relogin)
            ss = STORAGE_STATE if Path(STORAGE_STATE).exists() else None
            c = b.new_context(storage_state=ss) if ss else b.new_context()
            pg = c.new_page()
            pg.add_init_script(
                "Object.defineProperty(navigator, 'webdriver', {get: () => undefined});"
            )
            try:
                _ensure_logged_in_and_save_state(pg, c, STORAGE_STATE)
            except RuntimeError:
                # Sessão expirada ao iniciar → tenta re-login automático
                logger("⚠️  Sessão expirada ao iniciar — tentando login automático...")
                pg.close()
                b.close()
                if not auto_relogin(logger):
                    raise RuntimeError(
                        "Sessão expirada e login automático falhou. "
                        "Configure MANAGER_EMAIL e MANAGER_PASSWORD no Railway."
                    )
                # Recria browser com a nova sessão
                b = p.chromium.launch(headless=True, slow_mo=30)
                c = b.new_context(storage_state=STORAGE_STATE)
                pg = c.new_page()
                pg.add_init_script(
                    "Object.defineProperty(navigator, 'webdriver', {get: () => undefined});"
                )
                _ensure_logged_in_and_save_state(pg, c, STORAGE_STATE)
            return b, c, pg

        browser, context, page = _init_browser()

        def _is_cancelled() -> bool:
            if not job_id:
                return False
            from services.job_manager import get_job as _get_job
            j = _get_job(job_id)
            return j is not None and j.cancelled

        idx = 0
        while idx < len(dq_list):
            if _is_cancelled():
                logger("🛑 Extração cancelada pelo usuário.")
                break

            dq = dq_list[idx]
            questao = dq.block
            preview = (questao.enunciado[:90] + "...") if len(questao.enunciado) > 90 else questao.enunciado
            tipo_label = "RESERVA" if dq.is_reserva else "PRINCIPAL"
            logger(f"[{idx+1}/{total}] {tipo_label} {dq.label}: {preview}")

            try:
                match = _find_code_docx(page, questao, logger, job_id=job_id)
            except SessionExpiredError:
                logger("⚠️  Sessão expirada — tentando renovar automaticamente...")
                renewed = auto_relogin(logger)
                if not renewed:
                    logger("✗ Não foi possível renovar a sessão. Encerrando com resultados parciais.")
                    break
                try:
                    browser.close()
                except Exception:
                    pass
                browser, context, page = _init_browser()
                logger(f"✓ Sessão renovada. Retomando a partir de {dq.label}...")
                continue  # retry mesma questão sem avançar idx

            if match:
                found_count += 1
                logger(f"  ✓ ENCONTRADO: {match.code} (score={match.score_enunciado}%, {match.confianca}) [{found_count}/{total}]")
                entry = {"label": dq.label, "code": match.code}
                if dq.is_reserva:
                    reservas.append(entry)
                else:
                    principais.append(entry)
            else:
                logger(f"  ✗ NÃO ENCONTRADO [{found_count}/{total}]")
                entry = {"label": dq.label, "code": f"{dq.label}_NAO_ENCONTRADO"}
                if dq.is_reserva:
                    reservas.append(entry)
                else:
                    principais.append(entry)

            idx += 1

        browser.close()

    logger(f"CONCLUÍDO: {found_count}/{total} códigos encontrados.")
    logger(f"  Principais: {len(principais)} | Reservas: {len(reservas)}")

    return {
        "codes": principais + reservas,       # lista de {"label": "Q1", "code": "..."}
        "principais": principais,
        "reservas": reservas,
    }
